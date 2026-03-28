"""
CV export service — generates DOCX and PDF files from plain-text CV content.
Three templates: classic, modern, minimal.
"""
import io
import re
from typing import Literal

TemplateType = Literal["classic", "modern", "minimal"]


# ── CV text parser ────────────────────────────────────────────────────────────

def parse_cv(cv_text: str) -> list[tuple[str, str]]:
    """Parse plain-text CV into typed (tag, text) tuples.

    Tags: name | contact | section | bullet | body | blank
    Heuristics:
      - First non-empty line  → name
      - Second non-empty line that looks like contact info → contact
      - ALL-CAPS line (2–50 chars) that is not a number → section
      - Line starting with -, •, *, –, → → bullet
      - Everything else → body
    """
    lines = cv_text.split("\n")
    parsed: list[tuple[str, str]] = []
    first_non_empty = True
    expect_contact = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            parsed.append(("blank", ""))
            continue

        if first_non_empty:
            parsed.append(("name", stripped))
            first_non_empty = False
            expect_contact = True
            continue

        if expect_contact:
            expect_contact = False
            if re.search(r"[@|/]|linkedin|github|phone|tel|\+\d", stripped, re.IGNORECASE):
                parsed.append(("contact", stripped))
                continue

        # Section header: all-uppercase, reasonable length, not a digit-led line
        if (
            stripped == stripped.upper()
            and len(stripped) >= 3
            and len(stripped) <= 50
            and not stripped[0].isdigit()
            and re.search(r"[A-Z]", stripped)
        ):
            parsed.append(("section", stripped))
            continue

        # Bullet point
        if stripped[0] in ("-", "•", "*", "–", "→", "▪"):
            parsed.append(("bullet", stripped.lstrip("-•*–→▪ ")))
            continue

        parsed.append(("body", stripped))

    return parsed


# ── DOCX generator ────────────────────────────────────────────────────────────

def generate_docx(cv_text: str, template: TemplateType) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Page margins
    for sec in doc.sections:
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)

    # ── Template config ──
    if template == "classic":
        body_font = "Times New Roman"
        name_size, section_size, body_size = 18, 11, 10
        name_color    = RGBColor(0x1A, 0x20, 0x4B)  # deep navy
        section_color = RGBColor(0x1A, 0x20, 0x4B)
        body_color    = RGBColor(0x1A, 0x1A, 0x1A)
        contact_color = RGBColor(0x55, 0x55, 0x55)
        border_hex    = "1a204b"
        name_align    = WD_ALIGN_PARAGRAPH.CENTER
        contact_align = WD_ALIGN_PARAGRAPH.CENTER

    elif template == "modern":
        body_font = "Calibri"
        name_size, section_size, body_size = 20, 11, 10
        name_color    = RGBColor(0x1A, 0x36, 0x5D)  # deep blue
        section_color = RGBColor(0x2B, 0x6C, 0xB0)  # medium blue
        body_color    = RGBColor(0x2D, 0x37, 0x48)
        contact_color = RGBColor(0x4A, 0x55, 0x68)
        border_hex    = "2b6cb0"
        name_align    = WD_ALIGN_PARAGRAPH.LEFT
        contact_align = WD_ALIGN_PARAGRAPH.LEFT

    else:  # minimal
        body_font = "Calibri"
        name_size, section_size, body_size = 22, 9, 10
        name_color    = RGBColor(0x1A, 0x1A, 0x1A)
        section_color = RGBColor(0x71, 0x80, 0x96)  # slate gray
        body_color    = RGBColor(0x2D, 0x37, 0x48)
        contact_color = RGBColor(0x71, 0x80, 0x96)
        border_hex    = "cbd5e0"
        name_align    = WD_ALIGN_PARAGRAPH.LEFT
        contact_align = WD_ALIGN_PARAGRAPH.LEFT

    # ── Helpers ──
    def _style_run(run, size: int, color: RGBColor, bold: bool = False):
        run.font.name = body_font
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.bold = bold

    def _add_bottom_border(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), border_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    parsed = parse_cv(cv_text)

    for tag, text in parsed:
        if tag == "blank":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

        elif tag == "name":
            p = doc.add_paragraph()
            p.alignment = name_align
            p.paragraph_format.space_after = Pt(3)
            _style_run(p.add_run(text), name_size, name_color, bold=True)

        elif tag == "contact":
            p = doc.add_paragraph()
            p.alignment = contact_align
            p.paragraph_format.space_after = Pt(8)
            _style_run(p.add_run(text), body_size, contact_color)

        elif tag == "section":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            label = text.upper() if template == "minimal" else text
            _style_run(p.add_run(label), section_size, section_color, bold=True)
            _add_bottom_border(p)

        elif tag == "bullet":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(f"\u2022  {text}")
            _style_run(run, body_size, body_color)

        else:  # body
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _style_run(p.add_run(text), body_size, body_color)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF generator ─────────────────────────────────────────────────────────────

def generate_pdf(cv_text: str, template: TemplateType) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    # ── Template config ──
    if template == "classic":
        name_font    = "Times-Bold"
        body_font    = "Times-Roman"
        bold_font    = "Times-Bold"
        name_color   = HexColor("#1a204b")
        section_color= HexColor("#1a204b")
        body_color   = HexColor("#1a1a1a")
        contact_color= HexColor("#555555")
        line_color   = HexColor("#1a204b")
        name_size, section_size, body_size = 18, 11, 10
        name_align   = TA_CENTER
        contact_align= TA_CENTER
        uppercase_sections = False

    elif template == "modern":
        name_font    = "Helvetica-Bold"
        body_font    = "Helvetica"
        bold_font    = "Helvetica-Bold"
        name_color   = HexColor("#1a365d")
        section_color= HexColor("#2b6cb0")
        body_color   = HexColor("#2d3748")
        contact_color= HexColor("#4a5568")
        line_color   = HexColor("#2b6cb0")
        name_size, section_size, body_size = 20, 11, 10
        name_align   = TA_LEFT
        contact_align= TA_LEFT
        uppercase_sections = False

    else:  # minimal
        name_font    = "Helvetica"
        body_font    = "Helvetica"
        bold_font    = "Helvetica-Bold"
        name_color   = HexColor("#1a1a1a")
        section_color= HexColor("#718096")
        body_color   = HexColor("#2d3748")
        contact_color= HexColor("#718096")
        line_color   = HexColor("#cbd5e0")
        name_size, section_size, body_size = 22, 9, 10
        name_align   = TA_LEFT
        contact_align= TA_LEFT
        uppercase_sections = True

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
    )

    name_style = ParagraphStyle(
        "Name", fontName=name_font, fontSize=name_size,
        textColor=name_color, alignment=name_align, spaceAfter=4,
    )
    contact_style = ParagraphStyle(
        "Contact", fontName=body_font, fontSize=body_size,
        textColor=contact_color, alignment=contact_align, spaceAfter=10,
    )
    section_style = ParagraphStyle(
        "Section", fontName=bold_font, fontSize=section_size,
        textColor=section_color, spaceBefore=12, spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "Body", fontName=body_font, fontSize=body_size,
        textColor=body_color, spaceAfter=3, leading=14,
    )
    bullet_style = ParagraphStyle(
        "Bullet", fontName=body_font, fontSize=body_size,
        textColor=body_color, spaceAfter=2, leftIndent=14, leading=13,
    )

    story = []
    parsed = parse_cv(cv_text)

    def _safe(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    for tag, text in parsed:
        if tag == "blank":
            story.append(Spacer(1, 4))

        elif tag == "name":
            story.append(Paragraph(_safe(text), name_style))

        elif tag == "contact":
            story.append(Paragraph(_safe(text), contact_style))

        elif tag == "section":
            label = text.upper() if uppercase_sections else text
            story.append(Paragraph(_safe(label), section_style))
            story.append(HRFlowable(
                width="100%", thickness=0.75,
                color=line_color, spaceAfter=4,
            ))

        elif tag == "bullet":
            story.append(Paragraph(f"\u2022  {_safe(text)}", bullet_style))

        else:  # body
            story.append(Paragraph(_safe(text), body_style))

    doc.build(story)
    return buf.getvalue()
